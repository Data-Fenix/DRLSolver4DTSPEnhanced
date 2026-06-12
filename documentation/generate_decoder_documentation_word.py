"""
Script to generate DECODER_DOCUMENTATION.docx - Complete documentation
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_code_block(doc, code_text, language='python'):
    """Add a code block with monospace font"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add gray background (simulated with border)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_example_box(doc, title, content):
    """Add an example box"""
    p = doc.add_paragraph()
    p.add_run(f'Example: {title}').bold = True
    add_code_block(doc, content)

def create_decoder_documentation_word():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Decoder Architecture - Complete Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        'Overview',
        'Architecture Components',
        'Step-by-Step Process',
        'Code Implementation',
        'Workflows',
        'Examples',
        'Training and Evaluation'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== OVERVIEW ==========
    doc.add_heading('1. Overview', 1)
    
    doc.add_paragraph(
        'The decoder is the core component that solves the Dynamic Traveling Salesman Problem (DTSP) '
        'by sequentially selecting the next node to visit. It operates in an autoregressive manner, '
        'building a complete tour one node at a time.'
    )
    
    doc.add_heading('Main Inputs', 2)
    doc.add_paragraph('Node embeddings: Encoded representations of all nodes (first and previous)', style='List Bullet')
    doc.add_paragraph('Graph embeddings: Global graph representation', style='List Bullet')
    doc.add_paragraph('Visited states: Binary mask indicating which nodes have been visited', style='List Bullet')
    doc.add_paragraph('Current traffic condition: Time-dependent travel times between nodes', style='List Bullet')
    
    doc.add_heading('Main Outputs', 2)
    doc.add_paragraph('Selected node: Next node to visit', style='List Bullet')
    doc.add_paragraph('Log probabilities: For REINFORCE training', style='List Bullet')
    doc.add_paragraph('Tour cost: Total travel time', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== ARCHITECTURE COMPONENTS ==========
    doc.add_heading('2. Architecture Components', 1)
    
    # Step-MLP
    doc.add_heading('2.1 Step-MLP Enhancement (Optional)', 2)
    
    doc.add_paragraph(
        'The Step-MLP adds step-aware context to the decoder, helping it adapt its behavior based on '
        'the current progress in the tour.'
        'Two MLP layers can enhance attention and probability distribution:'
        '1. step-MLP: Adds a small bias (context_nudge) to the context vector'
        '2. temp-MLP: Scales attention logits by dividing by its output (temperature adjustment)'
    )
    
    doc.add_heading('Architecture', 3)
    
    p = doc.add_paragraph()
    p.add_run('Step-MLP:').bold = True
    add_code_block(doc, '''Input: 132 features (1 + 3 + 128)
  ↓
Linear(132 → 64) + ReLU
  ↓
Linear(64 → 64) + ReLU
  ↓
Linear(64 → 128)
  ↓
Output: context_nudge (128 dim)''')
    
    p = doc.add_paragraph()
    p.add_run('Temp-MLP:').bold = True
    add_code_block(doc, '''Input: 132 features (1 + 3 + 128)
  ↓
Linear(132 → 64) + ReLU
  ↓
Linear(64 → 1) + Sigmoid
  ↓
Output: temp_factor (0-1) → scaled to (0.5-2.5)''')
    
    doc.add_heading('Code Implementation', 3)
    
    add_code_block(doc, '''# From transformer.py lines 447-480

# Step-MLP initialization
if self.use_step_mlp:
    step_input_dim = (
        1 +  # k/N (current step ratio)
        3 +  # Last 3 visited nodes (binary encoding)
        embedding_dim  # Current state embedding
    )
    
    self.step_mlp = nn.Sequential(
        nn.Linear(step_input_dim, step_mlp_dim),
        nn.ReLU(),
        nn.Linear(step_mlp_dim, step_mlp_dim),
        nn.ReLU(),
        nn.Linear(step_mlp_dim, embedding_dim)  # Context nudge
    )
    
    # Temperature control MLP
    self.temp_mlp = nn.Sequential(
        nn.Linear(step_input_dim, step_mlp_dim),
        nn.ReLU(),
        nn.Linear(step_mlp_dim, 1),
        nn.Sigmoid()  # Output between 0-1
    )''')
    
    doc.add_heading('Step Features Extraction', 3)
    
    add_code_block(doc, '''# From transformer.py lines 501-570

def _get_step_features(self, state, embeddings):
    """Extract step features for MLP input"""
    batch_size = state.ids.size(0)
    graph_size = embeddings.size(1)
    
    # k/N ratio - current step normalized
    step_ratio = (state.i.float() / graph_size).unsqueeze(0).expand(batch_size, 1)
    
    # Last 3 visited nodes (binary encoding)
    last_3_nodes = torch.zeros(batch_size, 3, device=embeddings.device)
    if state.i.item() > 0:
        # Get indices of last visited nodes
        for i in range(min(3, state.i.item())):
            if state.i.item() - i - 1 >= 0:
                last_3_nodes[:, i] = 1.0  # Simple binary indicator
    
    # Current state embedding (mean of visited nodes)
    visited_mask = state.visited_  # Boolean mask of visited nodes
    visited_mask_expanded = visited_mask.squeeze(1).unsqueeze(-1)
    
    embeddings_clean = torch.where(torch.isnan(embeddings), 
                                   torch.zeros_like(embeddings), 
                                   embeddings)
    
    visited_embeddings = embeddings_clean * visited_mask_expanded.float()
    num_visited = visited_mask_expanded.sum(1).float()
    num_visited = torch.clamp(num_visited, min=1.0)
    
    state_embedding = visited_embeddings.sum(1) / num_visited
    
    # Concatenate all step features
    step_features = torch.cat([
        step_ratio,      # 1 dim
        last_3_nodes,    # 3 dim
        state_embedding  # 128 dim
    ], dim=1)  # Total: 132 dim
    
    return step_features''')
    
    add_example_box(doc, 'Step Features', '''# Example: Step 5 of 20 nodes, 3 nodes visited in last 3 steps
step_ratio = [0.25]  # 5/20
last_3_nodes = [1.0, 1.0, 1.0]  # All 3 positions indicate recent visits
state_embedding = [0.12, -0.45, 0.78, ...]  # 128-dim vector (mean of visited node embeddings)

step_features = [0.25, 1.0, 1.0, 1.0, 0.12, -0.45, 0.78, ...]  # 132 dim''')
    
    # Context Vector
    doc.add_heading('2.2 Context Vector Computation', 2)
    
    doc.add_paragraph(
        'The context vector combines information from multiple sources to guide node selection.'
    )
    
    doc.add_heading('Components', 3)
    doc.add_paragraph('1. Previous and First Node Embeddings (256 → 128 dim)', style='List Number')
    doc.add_paragraph('2. Visited States (128 dim)', style='List Number')
    doc.add_paragraph('3. Graph Embeddings (128 dim, static)', style='List Number')
    doc.add_paragraph('4. Current Traffic Condition (128 dim, dynamic)', style='List Number')
    
    doc.add_heading('Code Implementation', 3)
    
    add_code_block(doc, '''# From transformer.py lines 878-926

def _get_parallel_step_context(self, embeddings, state, mat, input):
    """
    Returns the context per step
    """
    b_s, i_s = embeddings.size(0), embeddings.size(1)
    _, ind = torch.max(input, dim=2)
    
    # Create coordinate matrices for (i,j) pairs
    xx_repeated = self.xx.repeat(b_s, 1, 1).view(b_s, i_s*i_s)
    yy_repeated = self.yy.repeat(b_s, 1, 1).view(b_s, i_s*i_s)
    
    # Get travel times for all pairs at current time
    getddd_result = mat.__getddd__(ind, xx_repeated, yy_repeated, state.lengths)
    reshaped = getddd_result.view(b_s, 1, i_s*i_s)
    
    # Project traffic condition to 128 dim
    current_traffic = self.project_traffic(reshaped)  # (batch_size, 1, 128)
    
    # Project visited states to 128 dim
    current_visit = self.project_visit(state.visited_.float())  # (batch_size, 1, 128)
    
    # Extract first and previous node embeddings
    ss = embeddings.gather(1, 
        torch.cat((state.first_a, state.prev_a), 1)[:, :, None]
        .expand(b_s, 2, embeddings.size(-1))
    )  # (batch_size, 2, 128)
    
    # Concatenate: [first+prev (256), traffic (128), visit (128)] = 512
    # Then projected to 128 in project_step_context
    return torch.cat((ss.view(b_s, 1, -1), current_traffic, current_visit), dim=2)''')
    
    # Traffic Condition
    doc.add_heading('Traffic Condition Computation', 3)
    
    doc.add_paragraph(
        'The traffic condition uses cubic spline interpolation to compute continuous travel times.'
    )
    
    add_code_block(doc, '''# From train.py lines 53-73 (DistanceMatrix.__getd__)

def __getd__(self, st, a, b, t):
    """
    Get travel time from node a to node b at time t using cubic spline
    """
    a = torch.gather(st, 1, a)
    b = torch.gather(st, 1, b)
    
    # Find time interval bounds
    tt = torch.floor(t * self.max_time_step) % self.max_time_step
    zz = (torch.floor(t * self.max_time_step) + 1) % self.max_time_step
    
    # Get indices for spline coefficients
    c = a.squeeze() * self.n_c * self.max_time_step + \\
        b.squeeze() * self.max_time_step + tt.squeeze().long()
    d = a.squeeze() * self.n_c * self.max_time_step + \\
        b.squeeze() * self.max_time_step + zz.squeeze().long()
    
    # Extract cubic spline coefficients
    a0 = torch.gather(self.mat, 0, c)  # Constant term
    a1 = torch.gather(self.m2, 0, c)    # Linear coefficient
    a2 = torch.gather(self.m3, 0, c)   # Quadratic coefficient
    a3 = torch.gather(self.m4, 0, c)    # Cubic coefficient
    b0 = torch.gather(self.mat, 0, d)   # Next interval constant
    
    # Compute fractional position within interval
    z = (t.squeeze() * self.max_time_step - 
         torch.floor(t.squeeze() * self.max_time_step)) / self.max_time_step
    z2 = z * z
    z3 = z2 * z
    
    # Cubic spline interpolation: a0 + a1*z + a2*z² + a3*z³
    res = a0 + a1 * z + a2 * z2 + a3 * z3
    
    # Clamp to reasonable bounds
    minres = (a0 + b0) * 0.05
    maxres = (a0 + b0) * 5
    res = torch.max(torch.cat((res.unsqueeze(-1), minres.unsqueeze(-1)), dim=-1), dim=-1)[0]
    res = torch.min(torch.cat((res.unsqueeze(-1), maxres.unsqueeze(-1)), dim=-1), dim=-1)[0]
    
    return res''')
    
    add_example_box(doc, 'Traffic Condition', '''# Example: Computing travel time from node 3 to node 7 at time 2.5
# max_time_step = 100, so time 2.5 maps to interval [250, 251]

# Extract coefficients for edge (3,7) at time 250
a0 = 1.2  # Constant
a1 = 0.05 # Linear
a2 = -0.001 # Quadratic  
a3 = 0.0001 # Cubic

# Fractional position: z = 0.5 (midpoint of interval)
z = 0.5
z2 = 0.25
z3 = 0.125

# Compute travel time
travel_time = 1.2 + 0.05*0.5 + (-0.001)*0.25 + 0.0001*0.125
            = 1.2 + 0.025 - 0.00025 + 0.0000125
            ≈ 1.225''')
    
    doc.add_page_break()
    
    # ========== MULTI-HEAD ATTENTION ==========
    doc.add_heading('2.3 Multi-Head Attention (MHA)', 2)
    
    doc.add_paragraph(
        'The MHA layer computes compatibility scores between the query and all nodes.'
    )
    
    doc.add_heading('Process', 3)
    doc.add_paragraph('1. Enrich node embeddings: Multiply by 3 → split into glimpse_key, glimpse_val, logit_key', style='List Number')
    doc.add_paragraph('2. Split into heads: 8 heads, each with 16 dimensions', style='List Number')
    doc.add_paragraph('3. Compute compatibility: compatibility = (query × glimpse_key^T) / √(key_size)', style='List Number')
    doc.add_paragraph('4. Mask visited nodes: Set compatibility to -∞', style='List Number')
    doc.add_paragraph('5. Apply softmax: Normalize compatibility scores', style='List Number')
    doc.add_paragraph('6. Weighted aggregation: weighted_value = softmax(compatibility) × glimpse_V', style='List Number')
    doc.add_paragraph('7. Concatenate heads: 8 heads × 16 dim = 128 dim', style='List Number')
    doc.add_paragraph('8. Final logits: logits = (glimpse × logit_K^T) / √(embed_dim)', style='List Number')
    
    doc.add_heading('Code Implementation', 3)
    
    add_code_block(doc, '''# From transformer.py lines 775-855

def _one_to_many_logits(self, query, glimpse_K, glimpse_V, logit_K, mask,
                        embeddings=None, state=None, mat=None, input=None):
    
    batch_size, num_steps, embed_dim = query.size()
    key_size = val_size = embed_dim // self.n_heads  # 128 / 8 = 16
    
    # Reshape query for multi-head attention
    # (batch_size, num_steps, n_heads, 1, key_size)
    glimpse_Q = query.view(batch_size, num_steps, self.n_heads, 1, key_size)\\
                     .permute(2, 0, 1, 3, 4)
    
    # Compute compatibility scores
    # (n_heads, batch_size, num_steps, graph_size)
    compatibility = torch.matmul(glimpse_Q, glimpse_K.transpose(-2, -1)) / \\
                   math.sqrt(glimpse_Q.size(-1))
    
    # Mask visited nodes
    if self.mask_inner:
        compatibility[mask[None, :, :, None, :].expand_as(compatibility)] = -math.inf
    
    # Apply softmax and compute weighted values
    # (n_heads, batch_size, num_steps, val_size)
    heads = torch.matmul(F.softmax(compatibility, dim=-1), glimpse_V)
    
    # Project and concatenate heads
    # (batch_size, num_steps, embedding_dim)
    glimpse = self.project_out(
        heads.permute(1, 2, 3, 0, 4).contiguous()
        .view(-1, num_steps, 1, self.n_heads * val_size)
    )
    
    # Compute final logits
    # (batch_size, num_steps, graph_size)
    logits = torch.matmul(glimpse, logit_K.transpose(-2, -1)).squeeze(-2) / \\
             math.sqrt(glimpse.size(-1))
    
    return logits, glimpse.squeeze(-2)''')
    
    add_example_box(doc, 'MHA Computation', '''# Example: 8 heads, 16 dim per head, 20 nodes

# Query: (batch=1, steps=1, embed_dim=128)
# glimpse_K: (8, batch=1, steps=1, nodes=20, key_size=16)

# Per head computation:
# Head 0: compatibility[0] = (query[0] × glimpse_K[0]^T) / √16
#         = (1×16 × 16×20) / 4 = (1×20)
#         After softmax: [0.05, 0.02, 0.15, ..., 0.01]  # 20 values sum to 1
#         weighted = softmax × glimpse_V[0] = (1×20 × 20×16) = (1×16)

# Concatenate all 8 heads: (1×16) × 8 = (1×128)
# Final logits: (1×128 × 128×20) / √128 = (1×20)''')
    
    doc.add_page_break()
    
    # ========== COST-AWARE GATING ==========
    doc.add_heading('2.4 Cost-Aware Gating (Optional)', 2)
    
    doc.add_paragraph(
        'Cost-aware gating adds heuristic-based bias to attention logits.'
    )
    
    doc.add_heading('Heuristic Types', 3)
    doc.add_paragraph('1. linear_time: Direct travel time from current position', style='List Number')
    doc.add_paragraph('2. nearest_neighbor: Distance to nearest unvisited node', style='List Number')
    doc.add_paragraph('3. nn_plus_one: Nearest neighbor with one-step lookahead', style='List Number')
    doc.add_paragraph('4. nnr: Probabilistic nearest neighbor random', style='List Number')
    
    doc.add_heading('Code Implementation', 3)
    
    add_code_block(doc, '''# From transformer.py lines 810-844

# NEW: Add Cost-Aware Gating (Soft Bias)
if self.use_cost_aware_gating and embeddings is not None:
    # Compute heuristic logits
    heuristic_logits = self.heuristic_computer.compute_heuristic_logits(
        embeddings, state, mat, input
    )  # Shape: (batch_size, graph_size)
    
    # Expand to match logits shape
    heuristic_logits = heuristic_logits.unsqueeze(1).expand(-1, num_steps, -1)
    
    # Apply nonlinear transformation if enabled
    if hasattr(self, 'transform'):
        if self.transform_type == 'piecewise':
            heuristic_logits = self.transform(heuristic_logits.unsqueeze(-1)).squeeze(-1)
        elif self.transform_type == 'exponential':
            heuristic_logits = self.exp_scale * torch.exp(heuristic_logits + self.exp_bias)
    
    # Add heuristic bias
    lambda_clipped = torch.clamp(self.lambda_heuristic, min=0.0, max=2.0)
    logits = logits + lambda_clipped * heuristic_logits''')
    
    add_example_box(doc, 'Cost-Aware Gating', '''# Example: Current position = node 5, 20 nodes total

# Base logits from attention: [0.2, -0.5, 1.3, 0.8, ..., -0.1]
# Heuristic (linear_time): [-1.2, -0.8, -2.1, -1.5, ..., -0.9]
# λ = 1.0

# Modified logits:
# logits[0] = 0.2 + 1.0 * (-1.2) = -1.0
# logits[1] = -0.5 + 1.0 * (-0.8) = -1.3
# logits[2] = 1.3 + 1.0 * (-2.1) = -0.8
# ...

# Node 2 (index 2) has shortest travel time, gets boosted''')
    
    doc.add_page_break()
    
    # ========== STEP-BY-STEP PROCESS ==========
    doc.add_heading('3. Step-by-Step Process', 1)
    
    doc.add_heading('Complete Decoding Loop', 2)
    
    add_code_block(doc, '''# From transformer.py lines 687-715

def _inner(self, input, embeddings, mat):
    outputs = []
    sequences = []
    
    # Initialize state
    state = StateTSP.initialize(input)
    state = state.addmask()  # Mark first node as visited
    
    # Precompute fixed attention data
    fixed = self._precompute(embeddings)
    
    # Decode until tour complete
    while not (state.all_finished()):
        # Step 1: Compute log probabilities
        log_p, mask = self._get_log_p(fixed, state, mat, input)
        
        # Step 2: Select next node
        selected = self._select_node(log_p.exp()[:, 0, :], mask[:, 0, :])
        
        # Step 3: Update state
        state = state.update(selected, mat, input)
        
        # Step 4: Store outputs
        outputs.append(log_p[:, 0, :])
        sequences.append(selected)
    
    # Return stacked outputs
    return torch.stack(outputs, 1), torch.stack(sequences, 1), state''')
    
    doc.add_heading('Flow Diagram', 2)
    
    flow_text = '''Initialize State
    ↓
[Loop: Until tour complete]
    ↓
Extract Step Features (if step-MLP enabled)
    ↓
Compute Context Vector
    ├─ Previous + First embeddings
    ├─ Visited states
    ├─ Graph embeddings
    └─ Current traffic condition
    ↓
Apply Step-MLP (optional)
    ├─ Add context_nudge
    └─ Get temp_adjustment
    ↓
Compute Query
    ↓
Multi-Head Attention
    ├─ Compute compatibility scores
    ├─ Mask visited nodes
    ├─ Apply softmax
    ├─ Weighted aggregation
    └─ Concatenate heads
    ↓
Compute Logits
    ↓
Apply Cost-Aware Gating (optional)
    └─ Add heuristic bias
    ↓
Apply Temperature Scaling
    ↓
Apply Softmax → Probabilities
    ↓
Select Next Node (Greedy/Sampling)
    ↓
Update State
    ├─ Update lengths
    ├─ Update prev_a
    ├─ Update visited
    └─ Increment step counter
    ↓
[End Loop]
    ↓
Return Tour and Log Probabilities'''
    
    add_code_block(doc, flow_text)
    
    doc.add_page_break()
    
    # ========== EXAMPLES ==========
    doc.add_heading('4. Examples', 1)
    
    doc.add_heading('Example 1: Complete Decoding Step', 2)
    
    example1 = '''# Scenario: 20-node TSP, currently at step 5, at node 3

# Step 1: Extract step features
step_ratio = 5 / 20 = 0.25
last_3_nodes = [1.0, 1.0, 0.0]  # Nodes visited in last 3 steps
state_embedding = mean(embeddings[visited_nodes])  # 128-dim vector

# Step 2: Compute context vector
prev_embedding = embeddings[:, 3, :]  # Node 3 embedding (128 dim)
first_embedding = embeddings[:, 0, :]  # First node embedding (128 dim)
visited_projection = project_visit(visited_mask)  # (128 dim)
graph_embedding = mean(embeddings)  # (128 dim)
traffic_condition = compute_traffic_at_time(current_time)  # (128 dim)

# Step 3: Apply step-MLP
context_nudge = step_mlp(step_features)  # (128 dim)
temp_adjustment = temp_mlp(step_features) * 2.0 + 0.5  # e.g., 1.2

# Step 4: Compute query
query = context_projected + project_step_context([prev, first, traffic, visit])
query = query + context_nudge

# Step 5: Multi-head attention
compatibility = (query × glimpse_key^T) / √16  # Per head
compatibility[mask] = -∞  # Mask visited nodes
attention_weights = softmax(compatibility)  # Per head
glimpse = concat([attention_weights × glimpse_val for each head])  # (128 dim)

# Step 6: Compute logits
logits = (glimpse × logit_key^T) / √128  # (20 dim)

# Step 7: Apply cost-aware gating
heuristic_logits = -travel_times  # Shorter = higher
logits = logits + 1.0 * heuristic_logits

# Step 8: Apply temperature and softmax
log_p = log_softmax(logits / (1.0 * 1.2))  # temp=1.0, adjustment=1.2
probs = exp(log_p)  # [0.05, 0.02, 0.35, 0.15, ...]

# Step 9: Select node
selected = argmax(probs)  # e.g., node 7

# Step 10: Update state
state.prev_a = 7
state.visited[7] = True
state.lengths += travel_time(3 → 7)
state.i += 1'''
    
    add_code_block(doc, example1)
    
    doc.add_heading('Example 2: Traffic Condition Computation', 2)
    
    example2 = '''# Scenario: Compute travel time from node 5 to node 12 at time 3.7
# max_time_step = 100, so time 3.7 maps to interval [370, 371]

# Step 1: Find time interval
tt = floor(3.7 * 100) % 100 = 70  # Lower bound
zz = (floor(3.7 * 100) + 1) % 100 = 71  # Upper bound

# Step 2: Get indices for edge (5, 12)
edge_idx = 5 * 20 + 12 = 112
c = 112 * 100 + 70 = 11270  # Index for time 70
d = 112 * 100 + 71 = 11271  # Index for time 71

# Step 3: Extract coefficients
a0 = mat[11270] = 1.5
a1 = m2[11270] = 0.03
a2 = m3[11270] = -0.001
a3 = m4[11270] = 0.0001

# Step 4: Compute fractional position
z = (3.7 * 100 - floor(3.7 * 100)) / 100 = 0.7
z2 = 0.49
z3 = 0.343

# Step 5: Interpolate
travel_time = 1.5 + 0.03*0.7 + (-0.001)*0.49 + 0.0001*0.343
            = 1.5 + 0.021 - 0.00049 + 0.0000343
            ≈ 1.5205

# Step 6: Clamp to bounds
min_bound = (1.5 + mat[11271]) * 0.05 = 0.15
max_bound = (1.5 + mat[11271]) * 5 = 15.0
travel_time = clamp(1.5205, 0.15, 15.0) = 1.5205'''
    
    add_code_block(doc, example2)
    
    doc.add_heading('Example 3: REINFORCE Loss Computation', 2)
    
    example3 = '''# Scenario: Batch of 4 instances

# Tour costs
costs = [12.5, 15.3, 11.8, 14.2]  # Total travel times

# Baseline values (from rollout baseline)
baseline = [13.0, 14.5, 12.0, 13.8]

# Log likelihoods (sum of log probs for selected actions)
log_likelihoods = [-2.3, -3.1, -2.0, -2.8]

# Compute advantages
advantages = costs - baseline
            = [12.5-13.0, 15.3-14.5, 11.8-12.0, 14.2-13.8]
            = [-0.5, 0.8, -0.2, 0.4]

# REINFORCE loss
reinforce_loss = mean(advantages * log_likelihoods)
               = mean([-0.5*-2.3, 0.8*-3.1, -0.2*-2.0, 0.4*-2.8])
               = mean([1.15, -2.48, 0.4, -1.12])
               = -0.51

# Interpretation:
# Instance 0: cost < baseline (better) → negative advantage → decreases loss
# Instance 1: cost > baseline (worse) → positive advantage → increases loss
# Instance 2: cost < baseline (better) → negative advantage → decreases loss
# Instance 3: cost > baseline (worse) → positive advantage → increases loss'''
    
    add_code_block(doc, example3)
    
    doc.add_page_break()
    
    # ========== TRAINING AND EVALUATION ==========
    doc.add_heading('5. Training and Evaluation', 1)
    
    doc.add_heading('Training Loop', 2)
    
    add_code_block(doc, '''# From train.py lines 226-260

def train_batch(mat, model, optimizer, baseline, epoch, batch_id, step, batch, tb_logger, opts):
    # Unwrap batch and move to device
    x, bl_val = baseline.unwrap_batch(batch)
    x = move_to(x, opts.device)
    bl_val = move_to(bl_val, opts.device) if bl_val is not None else None
    
    # Forward pass: decode tours
    cost, log_likelihood, _ = model(mat, x)
    
    # Evaluate baseline (if not provided)
    bl_val, bl_loss = baseline.eval(x, cost) if bl_val is None else (bl_val, 0)
    
    # Calculate REINFORCE loss
    reinforce_loss = ((cost - bl_val) * log_likelihood).mean()
    loss = reinforce_loss + bl_loss
    
    # Backward pass
    optimizer.zero_grad()
    loss.backward()
    grad_norms = clip_grad_norms(optimizer.param_groups, opts.max_grad_norm)
    optimizer.step()
    
    return loss, reinforce_loss''')
    
    doc.add_heading('Baseline Management', 2)
    
    add_code_block(doc, '''# From baselines.py lines 200-239

class RolloutBaseline(Baseline):
    def _update_model(self, model, epoch, dataset=None):
        # Save a copy of the model
        self.model = copy.deepcopy(model)
        
        # Generate or use provided dataset
        if dataset is None:
            self.dataset = TSPDataset(self.ci, size=self.opts.graph_size, 
                                     num_samples=self.opts.val_size, 
                                     distribution=self.opts.data_distribution)
        else:
            self.dataset = dataset
        
        # Evaluate baseline model on dataset
        print("Evaluating baseline model on evaluation dataset")
        self.bl_vals = rollout(self.mat, self.model, self.dataset, self.opts).cpu().numpy()
        self.mean = self.bl_vals.mean()
        self.epoch = epoch
    
    def epoch_callback(self, model, epoch):
        # Challenge baseline with current model
        candidate_vals = rollout(self.mat, model, self.dataset, self.opts).cpu().numpy()
        candidate_mean = candidate_vals.mean()
        
        # Update if significantly better (t-test)
        if candidate_mean - self.mean < 0:
            t, p = ttest_rel(candidate_vals, self.bl_vals)
            p_val = p / 2  # one-sided
            if p_val < self.opts.bl_alpha:
                print('Update baseline')
                self._update_model(model, epoch)''')
    
    doc.add_heading('First Epoch Baseline', 2)
    
    doc.add_paragraph(
        'For the first epoch, the baseline is initialized by:\n'
        '1. Rollout baseline: Evaluates the current model in greedy mode on a validation set\n'
        '2. Exponential baseline: Uses exponential moving average (starts with first batch costs)'
    )
    
    add_code_block(doc, '''# From baselines.py lines 90-120 (WarmupBaseline)

class WarmupBaseline(Baseline):
    def __init__(self, baseline, n_epochs=1, warmup_exp_beta=0.8):
        self.baseline = baseline
        self.n_epochs = n_epochs
        self.warmup_baseline = ExponentialBaseline(warmup_exp_beta)
    
    def unwrap_batch(self, batch):
        # During warmup, use exponential baseline
        if self.epoch < self.n_epochs:
            return self.warmup_baseline.unwrap_batch(batch)
        else:
            return self.baseline.unwrap_batch(batch)''')
    
    # Summary
    doc.add_page_break()
    doc.add_heading('Summary', 1)
    
    doc.add_paragraph(
        'The decoder architecture combines:\n'
        '1. Step-aware context via Step-MLP\n'
        '2. Dynamic traffic awareness via cubic spline interpolation\n'
        '3. Attention mechanisms via Multi-Head Attention\n'
        '4. Heuristic guidance via Cost-Aware Gating\n'
        '5. Reinforcement learning via REINFORCE with rollout baseline'
    )
    
    doc.add_paragraph(
        'This creates a powerful decoder capable of solving Dynamic TSP instances by adapting '
        'to both spatial and temporal constraints.'
    )
    
    # Save document
    doc.save('DECODER_DOCUMENTATION.docx')
    print("✓ Created DECODER_DOCUMENTATION.docx")

if __name__ == '__main__':
    create_decoder_documentation_word()