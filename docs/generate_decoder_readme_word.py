"""
Script to generate DECODER_README.docx from markdown content
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_decoder_readme_word():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Decoder Architecture - Quick Reference Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview Section
    doc.add_heading('Overview', 1)
    doc.add_paragraph(
        'The decoder sequentially selects the next node in a Dynamic Traveling Salesman Problem (DTSP) '
        'solution using attention mechanisms, step-aware MLPs, and cost-aware gating.'
    )
    
    # Key Components
    doc.add_heading('Key Components', 1)
    
    # Component 1
    doc.add_heading('1. Step-MLP Enhancement (Optional)', 2)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Adds context awareness and temperature control')
    p = doc.add_paragraph()
    p.add_run('Inputs: ').bold = True
    p.add_run('Step ratio (k/N), last 3 visited nodes, state embeddings')
    p = doc.add_paragraph()
    p.add_run('Outputs: ').bold = True
    p.add_run('\n  • context_nudge: Bias added to context vector (128 dim)\n  • temp_adjustment: Temperature scaling factor (0.5-2.5)')
    
    # Component 2
    doc.add_heading('2. Context Vector Assembly', 2)
    p = doc.add_paragraph()
    p.add_run('Components:').bold = True
    doc.add_paragraph('  • Previous + First node embeddings (256 → 128 after projection)', style='List Bullet')
    doc.add_paragraph('  • Visited states (128 dim)', style='List Bullet')
    doc.add_paragraph('  • Graph embeddings (128 dim, static)', style='List Bullet')
    doc.add_paragraph('  • Current traffic condition (128 dim, dynamic)', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Total: ').bold = True
    p.add_run('640 features → projected to 128 dimensions')
    
    # Component 3
    doc.add_heading('3. Multi-Head Attention (MHA)', 2)
    p = doc.add_paragraph()
    p.add_run('Heads: ').bold = True
    p.add_run('8 (default)')
    p = doc.add_paragraph()
    p.add_run('Process:').bold = True
    doc.add_paragraph('  • Compute compatibility scores', style='List Bullet')
    doc.add_paragraph('  • Apply masking for visited nodes', style='List Bullet')
    doc.add_paragraph('  • Weighted aggregation per head', style='List Bullet')
    doc.add_paragraph('  • Concatenate heads → 128 dim', style='List Bullet')
    
    # Component 4
    doc.add_heading('4. Cost-Aware Gating (Optional)', 2)
    p = doc.add_paragraph()
    p.add_run('Heuristics: ').bold = True
    p.add_run('linear_time, nearest_neighbor, nn_plus_one, nnr')
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run('logits = logits + λ × heuristic_logits')
    
    # Component 5
    doc.add_heading('5. Node Selection', 2)
    p = doc.add_paragraph()
    p.add_run('Methods: ').bold = True
    p.add_run('Greedy or Sampling')
    p = doc.add_paragraph()
    p.add_run('Temperature: ').bold = True
    p.add_run('log_p = log_softmax(logits / (temp × temp_adjustment))')
    
    # Component 6
    doc.add_heading('6. Reinforcement Learning', 2)
    p = doc.add_paragraph()
    p.add_run('Loss: ').bold = True
    p.add_run('reinforce_loss = ((cost - baseline) × log_likelihood).mean()')
    p = doc.add_paragraph()
    p.add_run('Baseline: ').bold = True
    p.add_run('Rollout baseline (evaluates best model so far)')
    
    # Training Flow
    doc.add_heading('Training Flow', 1)
    doc.add_paragraph('1. Initialize state', style='List Number')
    doc.add_paragraph('2. For each step until tour complete:', style='List Number')
    doc.add_paragraph('   • Compute context vector', style='List Bullet')
    doc.add_paragraph('   • Apply step-MLP (optional)', style='List Bullet')
    doc.add_paragraph('   • Compute attention logits', style='List Bullet')
    doc.add_paragraph('   • Apply cost-aware gating (optional)', style='List Bullet')
    doc.add_paragraph('   • Select next node', style='List Bullet')
    doc.add_paragraph('   • Update state', style='List Bullet')
    doc.add_paragraph('3. Compute REINFORCE loss', style='List Number')
    doc.add_paragraph('4. Backpropagate and update weights', style='List Number')
    
    # Key Files
    doc.add_heading('Key Files', 1)
    doc.add_paragraph('• transformer.py: Main decoder implementation', style='List Bullet')
    doc.add_paragraph('• heuristics.py: Heuristic computation methods', style='List Bullet')
    doc.add_paragraph('• train.py: Training loop and baseline management', style='List Bullet')
    doc.add_paragraph('• baselines.py: Baseline implementations', style='List Bullet')
    
    # Save document
    doc.save('DECODER_README.docx')
    print("✓ Created DECODER_README.docx")

if __name__ == '__main__':
    create_decoder_readme_word()